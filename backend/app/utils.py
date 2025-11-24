# backend/app/utils.py
# Enhanced Smart Text Extractor - with PDF Corruption Handling & Performance Fixes

import logging
from io import BytesIO
import pymupdf
from docx import Document
import re
from typing import Dict, List, Tuple, Optional
import asyncio
import unicodedata
from datetime import datetime
import json
from functools import lru_cache
import numpy as np  # ✅ نقل الاستيراد لأعلى الملف

logger = logging.getLogger(__name__)

# ============ Global OCR Reader (Singleton Pattern) ============

_ocr_reader = None
_ocr_initialized = False

@lru_cache(maxsize=1)
def get_global_ocr_reader(use_gpu: bool = False):
    """Get the OCR reader (global singleton)"""
    global _ocr_reader, _ocr_initialized
    
    if _ocr_initialized:
        return _ocr_reader
    
    _ocr_initialized = True
    
    try:
        import easyocr
        logger.info("🔄 Initializing global EasyOCR reader...")
        _ocr_reader = easyocr.Reader(
            ['en', 'ar'],
            gpu=use_gpu,
            verbose=False,
            quantize=True
        )
        logger.info("✅ EasyOCR reader initialized successfully")
        return _ocr_reader
    except Exception as e:
        logger.warning(f"⚠️ EasyOCR initialization failed: {e}")
        return None

# ============ Document Metadata ============

class DocumentMetadata:
    """Extracted document information"""
    def __init__(self):
        self.page_count = 0
        self.has_tables = False
        self.has_images = False
        self.is_scanned = False
        self.is_corrupted = False
        self.ocr_used = False
        self.language = "unknown"
        self.structure_type = "general"
        self.confidence = 0.0
        self.scanned_pages = []
        self.extraction_time_ms = 0.0
        self.corruption_details = ""

# ============ Smart Text Extractor ============

class SmartTextExtractor:
    """Enhanced smart text extractor with corruption handling"""
    
    CORRUPTION_INDICATORS = {
        'ﬁ', 'ﬂ', 'ﬃ', 'ﬀ', 'ﬄ', 'ﬅ', 'ﬆ',
        '¢', '¥', '£', '€', '¤',
        '†', '‡', '§', '¶',
        '•', '°', '‰', '‱',
        '©', '®', '™',
    }
    
    def __init__(self):
        self.metadata = DocumentMetadata()
        self._char_threshold = 100
        self.ocr_reader = get_global_ocr_reader(use_gpu=False)
    
    # ============ Corruption Detection & Fixing ============
    
    def _is_corrupted_text(self, text: str) -> bool:
        """Detect if the text is suspicious (Corrupted)"""
        if not text or len(text) < 50:
            return False
        
        sample = text[:10000]
        corruption_chars = sum(1 for char in sample if char in self.CORRUPTION_INDICATORS)
        corruption_ratio = corruption_chars / len(sample) if sample else 0
        
        unreadable_chars = sum(1 for char in sample if ord(char) > 65535 or unicodedata.category(char) == 'Cc')
        unreadable_ratio = unreadable_chars / len(sample) if sample else 0
        
        weird_spaces = sum(1 for char in sample if char in '\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u202F\u205F\u3000')
        weird_space_ratio = weird_spaces / len(sample) if sample else 0
        
        repeated_patterns = len(re.findall(r'(.)\1{4,}', sample))
        
        is_corrupted = (
            corruption_ratio > 0.05 or
            unreadable_ratio > 0.10 or
            weird_space_ratio > 0.08 or
            repeated_patterns > 3
        )
        
        if is_corrupted:
            logger.warning(
                f"⚠️ Text corruption detected: "
                f"corruption={corruption_ratio:.2%}, "
                f"unreadable={unreadable_ratio:.2%}, "
                f"weird_spaces={weird_space_ratio:.2%}, "
                f"repeated_patterns={repeated_patterns}"
            )
            self.metadata.is_corrupted = True
            self.metadata.corruption_details = (
                f"Corruption ratio: {corruption_ratio:.2%}, "
                f"Unreadable chars: {unreadable_ratio:.2%}"
            )
        
        return is_corrupted
    
    def _fix_corrupted_text(self, text: str) -> str:
        """Attempt to fix suspicious text"""
        if not text:
            return text
        
        try:
            ligature_map = {
                'ﬁ': 'fi', 'ﬂ': 'fl', 'ﬃ': 'ffi', 'ﬀ': 'ff',
                'ﬄ': 'ffl', 'ﬅ': 'ft', 'ﬆ': 'st'
            }
            for lig, replacement in ligature_map.items():
                text = text.replace(lig, replacement)
            
            weird_spaces = '\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u202F\u205F\u3000'
            for space in weird_spaces:
                text = text.replace(space, ' ')
            
            text = ''.join(char for char in text if ord(char) < 65535 and unicodedata.category(char)[0] != 'C')
            
            logger.info("✅ Text corruption fix applied")
            
        except Exception as e:
            logger.warning(f"⚠️ Corruption fix failed: {e}")
        
        return text
    
    # ============ OCR with Preprocessing ============
    
    def _preprocess_image_for_ocr(self, image):
        """Image optimization for OCR (speed-optimized)"""
        try:
            from PIL import Image, ImageEnhance, ImageFilter, ImageOps
            import cv2
            
            if image.mode != 'L':
                image = ImageOps.grayscale(image)
            
            img_array = np.array(image)
            
            img_array = cv2.adaptiveThreshold(
                img_array, 255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY, 11, 2
            )
            
            img_array = cv2.fastNlMeansDenoising(img_array, None, 10, 7, 21)
            image = Image.fromarray(img_array)
            
            width, height = image.size
            scale_factor = max(1, 3000 / max(width, height))
            if scale_factor > 1:
                new_size = (int(width * scale_factor), int(height * scale_factor))
                image = image.resize(new_size, Image.Resampling.BILINEAR)
            
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)
            
            image = image.filter(ImageFilter.SHARPEN)
            
            return image
        
        except Exception as e:
            logger.warning(f"⚠️ Image preprocessing failed: {e}")
            return image
    
    def _ocr_image_easyocr(self, image) -> str:
        """OCR using EasyOCR (optimized)"""
        if not self.ocr_reader:
            return ""
        
        try:
            image = self._preprocess_image_for_ocr(image)
            img_array = np.array(image)  # ✅ numpy مستورد بالفعل في الأعلى
            
            results = self.ocr_reader.readtext(
                img_array,
                detail=1,
                batch_size=8,
                paragraph=True,
                decoder='greedy',
                workers=0,
                beamWidth=5
            )
            
            if not results:
                return ""
            
            text_blocks = []
            total_confidence = 0.0
            valid_results = 0
            
            for item in results:
                if isinstance(item, (list, tuple)) and len(item) >= 3:
                    text, conf = item[1], item[2]
                    if conf > 0.3:
                        text_blocks.append(text)
                        total_confidence += conf
                        valid_results += 1
            
            if valid_results > 0:
                avg_page_conf = total_confidence / valid_results
                total_conf_pages = self.metadata.confidence * self.metadata.page_count
                self.metadata.confidence = (total_conf_pages + avg_page_conf) / (self.metadata.page_count + 1)
                
            return "\n".join(text_blocks)
        
        except Exception as e:
            logger.error(f"❌ EasyOCR failed: {e}")
            return ""
    
    def _extract_text_with_ocr_simple(self, doc, scanned_pages: List[int]) -> str:
        """OCR extraction from pages (optimized)"""
        if not scanned_pages:
            return ""
        
        logger.info(f"📄 Processing {len(scanned_pages)} scanned pages with OCR...")
        
        try:
            from PIL import Image
            
            results = {}
            for page_num in scanned_pages:
                page = doc[page_num]
                
                zoom_level = 1.5
                pix = page.get_pixmap(matrix=pymupdf.Matrix(zoom_level, zoom_level), alpha=False)
                
                image = Image.open(BytesIO(pix.tobytes("png")))
                page_text = self._ocr_image_easyocr(image)
                results[page_num] = page_text
                
                if page_text:
                    self.metadata.ocr_used = True
                    logger.debug(f"✓ Page {page_num + 1} OCR completed")
                
            all_text = []
            for page_num in sorted(results.keys()):
                if results[page_num]:
                    all_text.append(results[page_num])
            
            return "\n\n".join(all_text)
        
        except Exception as e:
            logger.error(f"❌ OCR extraction failed: {e}")
            return ""
    
    # ============ PDF/DOCX Extraction ============
    
    def _is_scanned_page(self, page_text: str) -> bool:
        """Detect a scanned page"""
        return len(page_text.strip()) < self._char_threshold
    
    def _is_scanned_pdf(self, doc) -> Tuple[bool, List[int]]:
        """Intelligently detect scanned pages"""
        scanned_pages = []
        total_pages = len(doc)
        
        sample_size = min(5, total_pages)
        step = max(1, total_pages // sample_size) if sample_size > 1 else 1
        sample_indices = list(range(0, total_pages, step))[:sample_size]
        
        for page_num in sample_indices:
            page = doc[page_num]
            text = page.get_text("text").strip()
            if self._is_scanned_page(text):
                scanned_pages.append(page_num)
        
        if scanned_pages and len(scanned_pages) / sample_size > 0.6:
            scanned_pages = []
            for page_num in range(total_pages):
                page = doc[page_num]
                text = page.get_text("text").strip()
                if self._is_scanned_page(text):
                    scanned_pages.append(page_num)
        
        return len(scanned_pages) > 0, scanned_pages
    
    def _detect_language(self, text: str) -> str:
        """Detect the language"""
        sample = text[:5000]
        sample_no_space = re.sub(r'\s+', '', sample)
        arabic_chars = len(re.findall(r'[\u0600-\u06FF]', sample_no_space))
        latin_chars = len(re.findall(r'[a-zA-Z]', sample_no_space))
        
        total = arabic_chars + latin_chars
        if total == 0:
            return "unknown"
        
        arabic_ratio = arabic_chars / total
        if arabic_ratio > 0.6:
            return "ar"
        elif arabic_ratio < 0.2:
            return "en"
        else:
            return "mixed"
    
    def _normalize_text(self, text: str) -> str:
        """Normalize the text"""
        text = ''.join(char for char in text 
                      if unicodedata.category(char)[0] != 'C' or char in '\n\t')
        
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'\s+([.,;:!?)])', r'\1', text)
        text = re.sub(r'([(])\s+', r'\1', text)
        
        return text.strip()
    
    def _extract_from_pdf(self, file_stream: BytesIO) -> str:
        """Extract from PDF (optimized with corruption handling)"""
        file_stream.seek(0)
        doc = pymupdf.open(stream=file_stream, filetype="pdf")
        self.metadata.page_count = len(doc)
        
        try:
            is_scanned, scanned_pages = self._is_scanned_pdf(doc)
            self.metadata.is_scanned = is_scanned
            self.metadata.scanned_pages = scanned_pages
            
            native_text_parts = []
            for page_num in range(len(doc)):
                if page_num not in scanned_pages:
                    page = doc[page_num]
                    page_text = page.get_text("text")
                    if page_text.strip():
                        native_text_parts.append(page_text)
                    
                    if page.find_tables():
                        self.metadata.has_tables = True
                    if page.get_images():
                        self.metadata.has_images = True
            
            native_text = "\n\n".join(native_text_parts)
            
            if self._is_corrupted_text(native_text):
                logger.warning("🔴 PDF corruption detected! Attempting to fix...")
                fixed_text = self._fix_corrupted_text(native_text)
                
                if self._is_corrupted_text(fixed_text):
                    logger.warning("📄 Corruption fix failed, checking OCR alternative...")
                    ocr_text_full = self._extract_text_with_ocr_simple(doc, list(range(len(doc))))
                    
                    if len(ocr_text_full) > len(fixed_text) * 0.7:
                        native_text = ocr_text_full
                        scanned_pages = []
                        self.metadata.ocr_used = True
                        logger.warning("✅ Switched to full OCR due to corruption")
                    else:
                        native_text = fixed_text
                else:
                    native_text = fixed_text
            
            ocr_text = ""
            if scanned_pages:
                if not self.ocr_reader:
                    logger.warning("⚠️ OCR Reader not available, skipping scanned pages")
                else:
                    ocr_text = self._extract_text_with_ocr_simple(doc, scanned_pages)
            
            all_text = [native_text] + ([ocr_text] if ocr_text else [])
            return "\n\n".join(all_text)
        
        finally:
            doc.close()
    
    def _extract_from_docx(self, file_stream: BytesIO) -> str:
        """Extract from DOCX"""
        file_stream.seek(0)
        doc = Document(file_stream)
        text_parts = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        
        for table in doc.tables:
            self.metadata.has_tables = True
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_text.append(row_text)
            if table_text:
                text_parts.append("\n".join(table_text))
        
        self.metadata.page_count = len(doc.sections)
        return "\n\n".join(text_parts)

# ============ Main Extraction Function ============

async def extract_text(file_stream: BytesIO, filename: str, use_ocr: bool = True, use_gpu: bool = False) -> str:
    """Enhanced smart text extraction with corruption handling (runs in a separate thread)"""
    start_time = datetime.now()
    
    try:
        raw_text = await asyncio.to_thread(_extract_sync_process, file_stream, filename, use_gpu)
        file_stream.seek(0)
        
        extractor = SmartTextExtractor()
        cleaned_text = extractor._normalize_text(raw_text)
        extraction_time = (datetime.now() - start_time).total_seconds() * 1000
        extractor.metadata.language = extractor._detect_language(cleaned_text)
        extractor.metadata.extraction_time_ms = extraction_time
        
        if len(cleaned_text) < 100:
            warning = "[Warning: Extracted text is too short. File may be corrupted.]"
            logger.warning(f"⚠️ {warning}")
            return warning
        
        if extractor.metadata.is_corrupted:
            logger.warning(f"⚠️ Corruption detected: {extractor.metadata.corruption_details}")
        
        logger.info(f"✅ Text extraction completed in {extraction_time:.0f}ms ({len(cleaned_text)} chars) | "
                    f"Language: {extractor.metadata.language}")
        
        return cleaned_text
    
    except Exception as e:
        logger.error(f"❌ Extraction failed: {str(e)}", exc_info=True)
        return f"[Error: Text extraction failed - {str(e)}]"

def _extract_sync_process(file_stream: BytesIO, filename: str, use_gpu: bool) -> str:
    """The synchronous function that performs the actual extraction"""
    file_stream.seek(0)
    extractor = SmartTextExtractor()
    extractor.ocr_reader = get_global_ocr_reader(use_gpu=use_gpu)
    
    if filename.lower().endswith(".pdf"):
        raw_text = extractor._extract_from_pdf(file_stream)
    elif filename.lower().endswith((".docx", ".doc")):
        raw_text = extractor._extract_from_docx(file_stream)
    else:
        return "[Error: Unsupported file type. Please upload PDF or DOCX files.]"
    
    return raw_text

def get_document_metadata(file_stream: BytesIO, filename: str) -> Dict:
    """Extract document metadata"""
    try:
        extractor = SmartTextExtractor()
        
        if filename.lower().endswith(".pdf"):
            file_stream.seek(0)
            doc = pymupdf.open(stream=file_stream, filetype="pdf")
            is_scanned, scanned_pages = extractor._is_scanned_pdf(doc)
            extractor.metadata.is_scanned = is_scanned
            extractor.metadata.scanned_pages = scanned_pages
            extractor.metadata.page_count = len(doc)
            
            for page in doc:
                if page.find_tables():
                    extractor.metadata.has_tables = True
                if page.get_images():
                    extractor.metadata.has_images = True
            
            doc.close()
        elif filename.lower().endswith((".docx", ".doc")):
            file_stream.seek(0)
            doc = Document(file_stream)
            extractor.metadata.page_count = len(doc.sections)
            
        return {
            'page_count': extractor.metadata.page_count,
            'has_tables': extractor.metadata.has_tables,
            'has_images': extractor.metadata.has_images,
            'is_scanned': extractor.metadata.is_scanned,
            'is_corrupted': extractor.metadata.is_corrupted,
            'scanned_pages_count': len(extractor.metadata.scanned_pages),
            'language': extractor.metadata.language,
            'structure_type': extractor.metadata.structure_type,
            'confidence': extractor.metadata.confidence,
            'extraction_time_ms': extractor.metadata.extraction_time_ms,
            'corruption_details': extractor.metadata.corruption_details,
            'ocr_used': extractor.metadata.ocr_used
        }
    
    except Exception as e:
        logger.error(f"Metadata extraction failed: {e}")
        return {}
